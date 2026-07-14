from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_figure(doc, image_path, caption):

    try:
        doc.add_picture(image_path, width=Inches(4.5))
    except Exception:
        return

    p = doc.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(caption)

    run.font.name = "Times New Roman"
    run.font.size = Pt(10)