"""
===============================================================================
IEEE DOCUMENT FORMATTER
===============================================================================

Responsibilities
----------------
✔ Create a new IEEE document
✔ Apply IEEE formatting rules
✔ Add Title
✔ Add Authors
✔ Add Affiliations
✔ Add Abstract
✔ Add Keywords
✔ Add Sections
✔ Add Tables
✔ Add Figures
✔ Add References

This module NEVER calls Ollama.

===============================================================================
"""

from docx import Document

from formatter.ieee_rules import (
    apply_page_layout,
    format_title,
    format_authors,
    format_affiliations,
    format_abstract_heading,
    format_abstract,
    format_keywords,
    format_heading1,
    format_body,
    format_table,
    format_table_caption,
    format_reference_heading,
    format_reference,
    format_figure_caption
)


def create_ieee_document(data, output_file):

    # -------------------------------------------------------------------------
    # Create Document
    # -------------------------------------------------------------------------

    doc = Document()

    apply_page_layout(doc)

    # -------------------------------------------------------------------------
    # Title
    # -------------------------------------------------------------------------

    if data.get("title"):

        p = doc.add_paragraph(data["title"])

        format_title(p)

    # -------------------------------------------------------------------------
    # Authors
    # -------------------------------------------------------------------------

    if data.get("authors"):

        authors = ", ".join(data["authors"])

        p = doc.add_paragraph(authors)

        format_authors(p)

    # -------------------------------------------------------------------------
    # Affiliations
    # -------------------------------------------------------------------------

    for aff in data.get("affiliations", []):

        p = doc.add_paragraph(aff)

        format_affiliations(p)

    # -------------------------------------------------------------------------
    # Abstract
    # -------------------------------------------------------------------------

    if data.get("abstract"):

        heading = doc.add_paragraph("Abstract")

        format_abstract_heading(heading)

        p = doc.add_paragraph(data["abstract"])

        format_abstract(p)

    # -------------------------------------------------------------------------
    # Keywords
    # -------------------------------------------------------------------------

    if data.get("keywords"):

        text = ", ".join(data["keywords"])

        p = doc.add_paragraph("Keywords: " + text)

        format_keywords(p)

    # -------------------------------------------------------------------------
    # Sections
    # -------------------------------------------------------------------------

    for section in data.get("sections", []):

        title = section.get("title", "")

        body = section.get("content", [])

        if title:

            p = doc.add_paragraph(title)

            format_heading1(p)

        for para in body:

            p = doc.add_paragraph(para)

            format_body(p)

    # -------------------------------------------------------------------------
    # Tables
    # -------------------------------------------------------------------------

    for table_data in data.get("tables", []):

        caption = table_data.get("caption", "")

        if caption:

            p = doc.add_paragraph(caption)

            format_table_caption(p)

        rows = table_data.get("rows", [])

        if rows:

            table = doc.add_table(
                rows=len(rows),
                cols=len(rows[0])
            )

            for i, row in enumerate(rows):

                for j, value in enumerate(row):

                    table.cell(i, j).text = value

            format_table(table)

    # -------------------------------------------------------------------------
    # Figures
    # -------------------------------------------------------------------------

    for fig in data.get("figures", []):

        image = fig.get("image", "")

        if image:

            try:
                doc.add_picture(image)
            except Exception:
                pass

        caption = fig.get("caption", "")

        if caption:

            p = doc.add_paragraph(caption)

            format_figure_caption(p)

    # -------------------------------------------------------------------------
    # References
    # -------------------------------------------------------------------------

    if data.get("references"):

        p = doc.add_paragraph("References")

        format_reference_heading(p)

        for ref in data["references"]:

            r = doc.add_paragraph(ref)

            format_reference(r)

    # -------------------------------------------------------------------------
    # Save
    # -------------------------------------------------------------------------

    doc.save(output_file)

    return output_file