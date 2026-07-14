from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_1(doc, text):
    p = doc.add_paragraph()

    run = p.add_run(text)

    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.bold = True

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading_2(doc, text):
    p = doc.add_paragraph()

    run = p.add_run(text)

    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.bold = True
    run.italic = True

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading_3(doc, text):
    p = doc.add_paragraph()

    run = p.add_run(text)

    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.bold = True

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT